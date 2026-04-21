"""
Build a polished, ATS-friendly Word CV for Rishav Singh (PM-focused).

Output: Rishav_Singh_PM_CV.docx (in repo root).
Run with:  python scripts/build_cv_docx.py
"""
from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────
# Styling helpers
# ─────────────────────────────────────────────────────────────

FONT = "Calibri"
# All-black colour scheme (no navy / grey accents).
ACCENT = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x00, 0x00, 0x00)
BODY = RGBColor(0x00, 0x00, 0x00)


def _set_run(run, *, size=9.5, bold=False, italic=False, color=BODY, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # also set East Asian font to keep Word happy on mixed locales
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def _add_para(doc, text="", *, size=9.5, bold=False, italic=False,
              color=BODY, align=None, space_after=1, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        run = p.add_run(text)
        _set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _add_h_rule(paragraph):
    """Add a thin horizontal rule under the given paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_heading(doc, text):
    p = _add_para(
        doc, text.upper(), size=10, bold=True, color=ACCENT,
        space_before=4, space_after=1,
    )
    _add_h_rule(p)
    return p


def _add_role_block(doc, *, company_dates: str, title: str, tagline: str,
                    client: str = "", client_label: str = "Client"):
    """Role header: company+dates, bold title, optional client line, italic tagline."""
    # company + dates
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(2)
    parts = company_dates.rsplit("·", 1)
    if len(parts) == 2:
        run = p.add_run(parts[0].rstrip())
        _set_run(run, size=9.5, bold=True)
        run2 = p.add_run("·" + parts[1])
        _set_run(run2, size=9.5)
    else:
        run = p.add_run(company_dates)
        _set_run(run, size=9.5, bold=True)

    # title
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run(title)
    _set_run(run, size=9.5, bold=True)

    # client / project line (own line, bold label)
    if client:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_after = Pt(0)
        run = p_c.add_run(f"{client_label}: ")
        _set_run(run, size=9.5, bold=True)
        run = p_c.add_run(client)
        _set_run(run, size=9.5)

    # tagline
    if tagline:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        run = p3.add_run(tagline)
        _set_run(run, size=9, italic=True)


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.4)
    # python-docx bullet style adds a default run; clear it and add our own
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(text)
    _set_run(run, size=9.5)
    return p


def _add_skills_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label + ": ")
    _set_run(run, size=9.5, bold=True, color=ACCENT)
    run2 = p.add_run(value)
    _set_run(run2, size=9.5)


# ─────────────────────────────────────────────────────────────
# Build the document
# ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Aggressive margins for strict one-page fit
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

    # ── Header: name + tagline + contact ──────────────────────
    name = _add_para(
        doc, "RISHAV KUMAR SINGH", size=16, bold=True, color=ACCENT,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0,
    )
    _add_para(
        doc, "Product Management  ·  AI Products  ·  Dublin",
        size=9.5, italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0,
    )
    contact = _add_para(
        doc,
        "+353 83 045 4123  ·  Stamp 1G  ·  rishav98sin@gmail.com  "
        "·  linkedin.com/in/rishav03sin  ·  Dublin, Ireland",
        size=8.5,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
    )
    _add_h_rule(contact)

    # ── Professional Summary ──────────────────────────────────
    _add_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.1
    bits = [
        ("Product builder with the technical depth to ship AI products end-to-end. ", False),
        ("Shipped two solo: ", False),
        ("ApplySmart AI", True),
        (", an agentic assistant that automates the full job-application workflow, and ", False),
        ("VoC Insight Hub", True),
        (", a customer-feedback synthesis tool for PMs. ", False),
        ("Four years at IBM and Accenture before that, delivering measurable wins on "
         "600K+ user platforms (40% efficiency gain, 25% user growth, 30% latency reduction). "
         "Targeting PM roles where technical depth, shipping speed, and shipped-product "
         "instincts matter. ", False),
        ("MSc Management (2.1), Trinity College Dublin.", False),
    ]
    for txt, is_bold in bits:
        run = p.add_run(txt)
        _set_run(run, size=9.5, bold=is_bold)

    # ── Featured Projects ─────────────────────────────────────
    _add_section_heading(doc, "Featured Projects")

    _add_role_block(
        doc,
        company_dates="ApplySmart AI · Apr 2026 – Present",
        title="Product Manager (solo)",
        tagline="An agentic AI product that automates the full job-application "
                "workflow for individual job seekers.",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("Tech Stack: ")
    _set_run(run, size=9, bold=True, color=ACCENT)
    run = p.add_run(
        "Python, LangGraph, LangChain, Groq LLaMA 3.3 70B, ChromaDB, "
        "Sentence-Transformers, Streamlit, PyMuPDF, ReportLab"
    )
    _set_run(run, size=9, italic=True)

    for b in [
        "Started from a problem I felt myself: every serious job application takes ~15 "
        "minutes of CV tweaking and cover-letter writing. Built a multi-agent system that "
        "does discovery, scoring, and tailoring end-to-end, cutting the same output to 2–3 "
        "minutes (~80% time saved per application).",
        "Scoped the agent flow as a supervisor pattern: one supervisor LLM routes between "
        "discovery, scoring, and generation workers across 5 job boards, with a 0–100 "
        "relevance scorer so only high-fit roles trigger the expensive tailoring step.",
        "Made trade-offs to keep the app usable on the Groq free tier: capped LLM calls "
        "per run, set a rate-limit ceiling so it aborts cleanly instead of hanging, and "
        "added a vector-retrieval layer that cut prompt size per match by ~50%.",
        "Built 9 guardrails against AI failure modes: a second LLM that grades every cover "
        "letter for fabrication and retries low scorers, input validation that blocks "
        "unusable CVs with clear reasons, and a preview-before-send mode so nothing "
        "accidental ever leaves the app. Owned every decision from problem to launch "
        "(scope, flows, UI, safety, deploy notes) and built crash-safe run snapshots into "
        "every session for post-launch iteration.",
    ]:
        _add_bullet(doc, b)

    _add_role_block(
        doc,
        company_dates="VoC Insight Hub · Jan – Mar 2026",
        title="Product Manager (solo)",
        tagline="A web app that helps PMs turn unstructured customer feedback into "
                "prioritised themes for stakeholder reviews.",
    )
    for label, body in [
        ("Problem:", " PMs sit on hundreds of free-text feedback rows with no scalable "
                     "way to surface what to actually do about them."),
        ("Solution:", " Upload feedback CSV, auto-cluster into themes, generate PM-ready "
                      "theme cards (problem, opportunity, success metric, confidence), "
                      "export to CSV or PPTX for stakeholder reviews."),
        ("Outcome:", " Live with real users. Full lifecycle ownership from discovery and "
                     "scoping through build and post-launch iteration."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(0)
        for r in list(p.runs):
            r.text = ""
        run = p.add_run(label)
        _set_run(run, size=9.5, bold=True)
        run = p.add_run(body)
        _set_run(run, size=9.5)

    # ── Professional Experience ───────────────────────────────
    _add_section_heading(doc, "Professional Experience")

    _add_role_block(
        doc,
        company_dates="IBM India Pvt. Ltd. · Aug 2023 – Jan 2024",
        title="Test Specialist – Performance & Resilience Management",
        client="Nationwide Building Society (UK)",
        tagline="Product-scope and performance work on a 600K+ member banking platform; "
                "partnered with the product team on scope, signals, and iteration.",
    )
    for b in [
        "Grew active users by 25% and lifted retention by 15% by translating performance "
        "signals into product recommendations, contributing to ~10% revenue growth on the platform.",
        "Translated vague user requests into clear scope documents and acceptance criteria "
        "for engineering, helping cut system latency by 30%.",
        "Ran 3+ product initiatives in parallel, keeping engineering and business stakeholders "
        "aligned and lifting delivery efficiency by 20%.",
    ]:
        _add_bullet(doc, b)

    _add_role_block(
        doc,
        company_dates="Accenture Solutions Pvt. Ltd. · Aug 2020 – Aug 2023",
        title="Performance Testing Engineer",
        client="Elevance Health (Fortune 25 US healthcare)",
        tagline="Owned scope, stakeholders, and iteration across 3+ applications from "
                "ideation through post-release. Recognised with Kudos and Spotlight "
                "awards for ownership and delivery.",
    )
    for b in [
        "Spotted an infrastructure bottleneck that was hurting users, built the business "
        "case and scope doc, secured stakeholder buy-in, and shipped changes that delivered "
        "40% efficiency gain and 2x platform capacity.",
        "Owned 3+ applications from idea through deployment and post-release; prioritised "
        "and unblocked 150+ user-impacting issues to keep platform quality and SLAs intact.",
        "Built a sprint-over-sprint performance dashboard that fed product decisions and "
        "surfaced regressions early, lifting baseline performance by 25%.",
        "Wrote monthly production-analysis briefs for leadership, turning dashboards and "
        "incidents into clear \u201chere\u2019s what to do next\u201d recommendations.",
    ]:
        _add_bullet(doc, b)

    # ── Education ───────────────────────────────────────────
    _add_section_heading(doc, "Education")
    for year, degree, inst, grade in [
        ("Sep 2024 – Oct 2025", "MSc Management", "Trinity College Dublin", "2.1"),
        ("Jul 2016 – May 2020", "B.Tech Mechanical Engineering",
         "KIIT University, Bhubaneswar", "7.57"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{degree}, ")
        _set_run(run, size=9.5, bold=True)
        run = p.add_run(f"{inst} · {year} · Grade: {grade}")
        _set_run(run, size=9.5)

    # ── Skills ────────────────────────────────────────────────
    _add_section_heading(doc, "Skills")
    _add_skills_line(
        doc, "Product",
        "Product roadmapping, PRD and scope-doc writing, Agile/Scrum, "
        "RICE/MoSCoW prioritisation, sprint planning, user stories, acceptance criteria, "
        "OKRs, product discovery, user flow design, A/B testing, iterative delivery",
    )
    _add_skills_line(
        doc, "Delivery",
        "Stakeholder management, cross-functional leadership, requirements elicitation, "
        "risk management, data-driven problem solving, executive communication, "
        "strategic thinking",
    )
    _add_skills_line(
        doc, "AI",
        "Large Language Models, prompt engineering, agentic AI systems, multi-agent "
        "architecture, Groq API, RAG, vector databases, LLM-based scoring, workflow "
        "automation, Mixpanel, Cursor",
    )
    _add_skills_line(
        doc, "Tools",
        "JIRA, Notion, Miro, SQL, Python, Java, Streamlit, LangGraph, LangChain, "
        "ChromaDB, APIs, DynaTrace, Splunk, UNIX/Linux, LoadRunner/VuGen, "
        "Oracle Enterprise Manager, Microsoft Office",
    )

    # ── Save ──────────────────────────────────────────────────
    out_path = os.path.abspath("Rishav_Singh_PM_CV.docx")
    doc.save(out_path)
    print(f"\u2713 Wrote: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
