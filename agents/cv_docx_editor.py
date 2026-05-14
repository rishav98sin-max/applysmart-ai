"""
agents.cv_docx_editor
=====================

Apply a tailor-produced diff to a DOCX file (May 2026 / DOCX path).

The diff shape (from `cv_diff_tailor.tailor_cv_diff`) is:

    {
        "summary": "new summary text",          # "" or absent → keep original
        "bullets": {
            "Role Header Match Key": [
                {"i": 0, "text": "new bullet text"},   # rewrite bullet 0
                {"i": 1, "text": None},                # keep bullet 1
                {"i": 2},                              # keep bullet 2 (text omitted)
                ...
            ],
            ...
        },
        "skills_order": [],   # always empty per current policy
    }

Anchors flow through the outline as `_anchor` (per role / per bullet) and
`_summary_anchors` (list — one summary may span multiple paragraphs).
We use these to relocate paragraphs in the on-disk DOCX without re-parsing.

Run preservation policy
-----------------------
A paragraph can contain multiple runs (each with its own bold/italic/font/
size/colour). When we rewrite a bullet's text we:

  1.  Capture the FIRST non-empty run's formatting (the "primary" style).
  2.  Replace all runs with a single run carrying the new text + primary
      formatting.

This loses mid-paragraph formatting variations (e.g. a bolded sub-phrase
inside a bullet) but preserves what matters for CV rendering: paragraph
style (List Bullet / List Paragraph), alignment, font family, base font
size, and colour. The trade-off matches the PDF in-place editor's
behaviour, which also loses sub-bullet style variation.

Summary handling
----------------
A summary often spans multiple paragraphs (e.g. one prose sentence per
line, or a header line followed by a body). We collapse the new summary
text into the FIRST summary paragraph and CLEAR the rest. This avoids
length-mismatch artefacts (e.g. shipping a 3-paragraph summary block
when the new summary is one sentence).

Bullet reorder support
----------------------
v1 keeps bullets in their original document order — same as the PDF
in-place editor. The reorder field is read but only the indices' order
of TEXT REWRITES is honoured (paragraph order in the DOCX stays put).
This matches the current product behaviour where reorder hasn't
demonstrated quality gains over targeted rewrites.

Never raises. On any failure returns `(False, reason_string)` so the
caller can fall back to the rebuild path.
"""

from __future__ import annotations

import os
import re
from typing import Any, Dict, List, Optional, Tuple


# Bullet glyphs that may sit at the start of a paragraph's text (vs being
# rendered by Word's list numbering machinery). When the original
# paragraph uses an in-text glyph (common in pdf2docx-converted CVs), we
# must preserve it on rewrite — otherwise the parser doesn't recognise
# the rewritten paragraph as a bullet on the next pass.
_LEADING_GLYPH_RX = re.compile(
    r"^(\s*[\u2022\u00b7\u25aa\u25cb\u25a0\u2043\u2219\u25b8\u25b6]\s*"
    r"|\s*[\-\*]\s+)"
)


# ─────────────────────────────────────────────────────────────
# Run / paragraph mutation helpers
# ─────────────────────────────────────────────────────────────

def _first_non_empty_run_index(runs: List[Any]) -> int:
    """
    Return the index into `runs` of the first run whose text is non-empty.
    Falls back to 0 if every run is blank. Caller MUST pass a single
    snapshot of `list(paragraph.runs)` and reuse the returned index against
    that same list — `paragraph.runs` returns fresh proxy objects on every
    access, so identity checks across two calls do NOT match (the bug that
    made every rewrite silently blank out the paragraph in the May 13
    smoke test).
    """
    for i, r in enumerate(runs):
        if (r.text or "").strip():
            return i
    return 0


def _replace_paragraph_text(paragraph: Any, new_text: str) -> None:
    """
    Replace the paragraph's visible text while preserving paragraph-level
    formatting (style, alignment, indent) and the FIRST run's font/bold/
    italic/colour properties.

    Implementation notes:
      - We don't simply set `paragraph.text = "..."` because that
        path destroys every run, including their font properties.
      - We don't delete runs naively because some runs carry properties
        (e.g. tab stops) that other runs reference.
      - Instead: rewrite the first non-empty run's text, then blank out
        every subsequent run's text. This leaves the run XML intact while
        showing only the new text.
      - We work with a SINGLE snapshot of `paragraph.runs` and index by
        position. `paragraph.runs` returns fresh proxy objects each call,
        so `proxy_a is proxy_b` is False even for the same underlying XML.
    """
    if paragraph is None:
        return
    runs = list(paragraph.runs)
    if not runs:
        # Empty paragraph — add a single run with the new text.
        paragraph.add_run(new_text)
        return
    # Preserve any leading bullet glyph from the original text. This
    # matters for paragraphs where the bullet rendering comes from an
    # in-text glyph rather than Word's list numbering (typical of
    # pdf2docx output: "•Started from a problem…"). Stripping the glyph
    # would make the parser treat the rewritten paragraph as prose on
    # the next pass, losing its bullet identity entirely.
    original_text = paragraph.text or ""
    glyph_match = _LEADING_GLYPH_RX.match(original_text)
    if glyph_match and not _LEADING_GLYPH_RX.match(new_text):
        new_text = glyph_match.group(1) + new_text

    primary_idx = _first_non_empty_run_index(runs)
    runs[primary_idx].text = new_text
    for i, r in enumerate(runs):
        if i == primary_idx:
            continue
        r.text = ""


def _blank_paragraph(paragraph: Any) -> None:
    """Set every run in the paragraph to empty text (visible result: blank line)."""
    if paragraph is None:
        return
    for r in paragraph.runs:
        r.text = ""


def _index_paragraphs(doc: Any) -> Dict[int, Any]:
    """
    Build {global_paragraph_index → paragraph_object} matching the same
    walk order used by `cv_docx_parser._iter_body_paragraphs` — i.e.
    TRUE document order with table-cell paragraphs interleaved at the
    position the table appears in the body.

    The two walks MUST be in lock-step. Any divergence (e.g. one walks
    paragraphs-then-tables, the other walks document-order) silently
    misaligns anchors and the editor rewrites the wrong paragraph.
    """
    from docx.oxml.ns import qn   # late import

    try:
        from docx.text.paragraph import Paragraph as _Paragraph
        from docx.table import Table as _Table
    except Exception:
        _Paragraph = _Table = None   # type: ignore

    P_TAG  = qn("w:p")
    TBL_TAG = qn("w:tbl")

    out: Dict[int, Any] = {}
    idx = 0

    def _walk(parent_element: Any, container_obj: Any) -> None:
        nonlocal idx
        for child in parent_element.iterchildren():
            tag = child.tag
            if tag == P_TAG:
                if _Paragraph is not None:
                    p = _Paragraph(child, container_obj)
                else:
                    p = doc.paragraphs[0].__class__(child, container_obj)
                out[idx] = p
                idx += 1
            elif tag == TBL_TAG and _Table is not None:
                tbl = _Table(child, container_obj)
                for row in tbl.rows:
                    for cell in row.cells:
                        _walk(cell._tc, cell)

    _walk(doc.element.body, doc)
    return out


# ─────────────────────────────────────────────────────────────
# Diff application
# ─────────────────────────────────────────────────────────────

def apply_diff_to_docx(
    docx_path:   str,
    diff:        Dict[str, Any],
    outline:     Dict[str, Any],
    output_path: str,
) -> Tuple[bool, str]:
    """
    Open `docx_path`, apply `diff` (matched against `outline`'s anchors),
    save to `output_path`. Returns `(True, "")` on success or
    `(False, reason)` on any failure (file not found, no python-docx,
    write error, etc.).

    The function is intentionally fault-tolerant:
      - Missing role anchors (e.g. tailor renamed the role and our match
        already failed upstream) are skipped, not fatal.
      - Out-of-range bullet indices are skipped.
      - A `None` / missing `text` on a bullet means "keep original" and
        leaves the paragraph untouched.

    Output is byte-true to the input format (DOCX), so the caller can
    then route through `cv_docx_to_pdf.render_pdf` to produce the final
    PDF.
    """
    if not docx_path or not os.path.exists(docx_path):
        return False, f"DOCX file not found: {docx_path!r}"

    try:
        import docx as _docx_lib
    except Exception as e:
        return False, f"python-docx unavailable: {type(e).__name__}: {e}"

    try:
        doc = _docx_lib.Document(docx_path)
    except Exception as e:
        return False, f"Failed to open DOCX: {type(e).__name__}: {e}"

    paragraphs = _index_paragraphs(doc)
    if not paragraphs:
        return False, "DOCX has no paragraphs"

    summary_text = (diff.get("summary") or "").strip()
    summary_anchors: List[int] = outline.get("_summary_anchors") or []

    edits_applied: int = 0
    edits_skipped: int = 0

    # ── 1. Summary ──────────────────────────────────────────
    # If the diff carries a non-empty summary AND we know which paragraphs
    # the original summary occupies, rewrite the first one with the new
    # text and clear any continuation paragraphs.
    if summary_text and summary_anchors:
        first_idx = summary_anchors[0]
        first_para = paragraphs.get(first_idx)
        if first_para is not None:
            _replace_paragraph_text(first_para, summary_text)
            edits_applied += 1
        # Clear any continuation paragraphs (rare — most summaries are one
        # paragraph) so we don't ship stale prose alongside the new summary.
        for cont_idx in summary_anchors[1:]:
            cont_para = paragraphs.get(cont_idx)
            if cont_para is not None:
                _blank_paragraph(cont_para)

    # ── 2. Bullets ──────────────────────────────────────────
    # The diff's `bullets` dict is keyed by role header. We've already
    # done the fuzzy role-key matching in `_sanitise_diff` (cv_diff_tailor),
    # so the keys here align 1:1 with `outline["roles"][k]["header"]`.
    # Iterate the outline's roles and look up the matching diff entry.
    bullets_diff: Dict[str, Any] = diff.get("bullets") or {}
    if isinstance(bullets_diff, dict):
        # Build a lowercase header → diff-entries map for case-tolerant lookup.
        diff_by_header_l: Dict[str, List[Dict[str, Any]]] = {
            str(k).strip().lower(): v
            for k, v in bullets_diff.items()
            if isinstance(v, list)
        }
        for role in outline.get("roles", []):
            header = (role.get("header") or "").strip()
            header_l = header.lower()
            entries = diff_by_header_l.get(header_l)
            if not entries:
                continue
            role_bullets = role.get("bullets") or []
            n_bullets = len(role_bullets)

            # Run 19 audit fix: pre-pass to group megabullet siblings.
            # When pdf2docx merged multiple bullets into one paragraph, the
            # parser split them into atomic bullets that share an _anchor.
            # The editor needs to write back the joined result (mix of new
            # rewrites for changed bullets + original text for kept ones)
            # as a single paragraph so the layout doesn't break.
            megabullet_groups: Dict[int, List[Dict[str, Any]]] = {}
            for i, b in enumerate(role_bullets):
                if not isinstance(b, dict):
                    continue
                if b.get("_megabullet_count", 1) > 1:
                    anchor_key = b.get("_anchor")
                    if isinstance(anchor_key, int):
                        megabullet_groups.setdefault(anchor_key, []).append(
                            {"role_index": i, "bullet": b}
                        )
            mega_anchors = set(megabullet_groups.keys())
            # Map role_index → final text after applying diff
            mega_resolved: Dict[int, Dict[int, str]] = {
                a: {} for a in mega_anchors
            }

            for entry in entries:
                if not isinstance(entry, dict):
                    continue
                try:
                    i = int(entry.get("i"))
                except (TypeError, ValueError):
                    continue
                if i < 0 or i >= n_bullets:
                    edits_skipped += 1
                    continue
                new_text = entry.get("text")
                if not isinstance(new_text, str) or not new_text.strip():
                    # Caller marked this bullet "keep original" — no-op
                    # for atomic bullets, but for megabullet siblings we
                    # still need to record the ORIGINAL text so the joined
                    # output is complete.
                    bullet_obj = role_bullets[i]
                    if isinstance(bullet_obj, dict):
                        anchor_k = bullet_obj.get("_anchor")
                        if anchor_k in mega_anchors:
                            mega_resolved[anchor_k][i] = (
                                bullet_obj.get("text") or ""
                            ).strip()
                    continue
                # Locate the actual DOCX paragraph for bullet i.
                bullet_obj = role_bullets[i]
                if not isinstance(bullet_obj, dict):
                    edits_skipped += 1
                    continue
                anchor = bullet_obj.get("_anchor")
                if not isinstance(anchor, int):
                    edits_skipped += 1
                    continue
                target_para = paragraphs.get(anchor)
                if target_para is None:
                    edits_skipped += 1
                    continue

                # Megabullet path: stash the rewrite, apply later after
                # collecting all siblings.
                if anchor in mega_anchors:
                    mega_resolved[anchor][i] = new_text.strip()
                    edits_applied += 1
                    continue

                # Atomic bullet path (unchanged).
                _replace_paragraph_text(target_para, new_text.strip())
                # Blank any continuation paragraphs the parser folded into
                # this bullet's text. Without this the rendered PDF shows
                # the NEW bullet followed by the ORIGINAL wrap-line text
                # below it (e.g. "[New rewrite] ... [original wrap-text]"
                # leaks through). See cv_docx_parser._continuation_anchors.
                #
                # Run-17 audit fix #27: the parser's continuation detection
                # is permissive — any prose paragraph after a bullet can be
                # misclassified as a wrap-continuation. Blanking such a
                # paragraph silently deletes legitimate CV content. We now
                # only blank continuations that LOOK like genuine wrap text:
                #   - short (≤ 60% of typical line length, i.e. ≤72 chars)
                #   - OR lowercase-starting (a wrap line doesn't start a
                #     new sentence)
                # Any continuation that looks like an independent paragraph
                # (long, capitalised-starting) is left in place and logged.
                cont_anchors = bullet_obj.get("_continuation_anchors") or []
                for cont_idx in cont_anchors:
                    cont_para = paragraphs.get(cont_idx)
                    if cont_para is None:
                        continue
                    cont_text = (cont_para.text or "").strip()
                    if not cont_text:
                        # Already empty — safe to skip.
                        continue
                    looks_like_wrap = (
                        len(cont_text) <= 72
                        or (cont_text[:1].islower() if cont_text else False)
                        or cont_text.startswith(("•", "-", "·", "*"))
                    )
                    if looks_like_wrap:
                        _blank_paragraph(cont_para)
                    else:
                        print(
                            f"   🛡️  cv_docx_editor: kept continuation @ "
                            f"anchor={cont_idx} (looks independent, "
                            f"len={len(cont_text)}, preview="
                            f"{cont_text[:60]!r})"
                        )
                edits_applied += 1

            # Run 19 audit fix: write back the megabullet groups now that
            # we've collected all sibling rewrites. Each group ends up as
            # a single paragraph with the bullets joined by newlines and
            # bullet glyphs (preserves the visual list structure).
            for anchor_k, role_index_to_text in mega_resolved.items():
                if not role_index_to_text:
                    continue
                target_para = paragraphs.get(anchor_k)
                if target_para is None:
                    continue
                # Build the joined text in original sibling order. For any
                # sibling that wasn't in the diff (no rewrite, no explicit
                # keep), use its parsed original text so nothing gets lost.
                sibling_entries = megabullet_groups[anchor_k]
                # Sort siblings by sub-index (preserves visual order)
                sibling_entries.sort(
                    key=lambda e: e["bullet"].get("_megabullet_subidx", 0)
                )
                joined_lines: List[str] = []
                for s_entry in sibling_entries:
                    s_i = s_entry["role_index"]
                    s_bullet = s_entry["bullet"]
                    if s_i in role_index_to_text:
                        joined_lines.append(role_index_to_text[s_i])
                    else:
                        joined_lines.append((s_bullet.get("text") or "").strip())
                joined_lines = [ln for ln in joined_lines if ln]
                if joined_lines:
                    # Join with newline + bullet glyph so LibreOffice renders
                    # them as visual bullets within the same paragraph.
                    glyph_sep = "\n• "
                    joined_text = "• " + glyph_sep.join(joined_lines)
                    _replace_paragraph_text(target_para, joined_text)
                    print(
                        f"   🔗 cv_docx_editor: wrote megabullet @ "
                        f"anchor={anchor_k} ({len(joined_lines)} bullets "
                        f"joined, {len(joined_text)} chars)"
                    )

    # ── 3. Save ─────────────────────────────────────────────
    try:
        # Ensure output directory exists.
        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)
        doc.save(output_path)
    except Exception as e:
        return False, f"Failed to save DOCX: {type(e).__name__}: {e}"

    print(
        f"   ✏️  cv_docx_editor: applied={edits_applied} skipped={edits_skipped} "
        f"→ {os.path.basename(output_path)}"
    )
    return True, ""
