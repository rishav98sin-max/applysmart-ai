"""
agents.cv_pdf_to_docx
=====================

Convert an uploaded PDF CV to DOCX (May 2026 / DOCX path) so the rest of
the tailoring pipeline can work on a flow-based document instead of
fighting PDF coordinate geometry.

Convertibility score
--------------------
`pdf2docx` works well on plain single-column, single-font CVs and degrades
on:

  - Multi-column layouts (Canva, NovoResume, ATS-template-style sidebars)
  - Custom fonts / icon glyphs
  - Coloured backgrounds, banners, or photo headers
  - Tables-as-layout designs

For those, the conversion produces a DOCX whose structure looks fine to
python-docx but whose visual rendering loses the original styling. We
catch this BEFORE editing by computing a `convertibility_score` (0-100)
on the converted DOCX:

  +60 : produces at least the same paragraph / bullet count as the PDF
  +20 : section headers detected after conversion (≥2 known section labels)
  +10 : no multi-column tables in the DOCX body (no sidebar / two-col layout)
  +10 : converted DOCX < 3× the original PDF byte size (sanity check)

Default acceptance threshold: 60. Below that, the router falls back to
the existing WeasyPrint rebuild path (designer-CV friendly).

Never raises. On any failure returns
`{"ok": False, "score": 0, "reason": "..."}`.
"""

from __future__ import annotations

import os
import re
from typing import Any, Dict, Optional


# Reuse the same section regex used by the validator so a converted DOCX
# is judged by the same rubric the human-uploaded DOCX is.
_SECTION_RX = re.compile(
    r"\b("
    r"summary|profile|objective|about(\s+me)?|"
    r"experience|professional\s+experience|work\s+experience|employment\s*history|"
    r"education|academic(\s+achievements?|\s+background)?|"
    r"skills?|technical\s+skills?|core\s+competenc(?:y|ies)|"
    r"projects?|featured\s+projects?|portfolio|"
    r"certifications?|awards?|publications?|languages?"
    r")\b",
    re.IGNORECASE,
)

# Acceptance threshold. Below this we fall back to the rebuild path.
# Tunable via env if we discover deployments where the pdf2docx fidelity
# is unexpectedly low or high.
CONVERTIBILITY_THRESHOLD = int(
    os.getenv("DOCX_CONVERTIBILITY_THRESHOLD", "60")
)


# ─────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────

def _measure_pdf(pdf_path: str) -> Dict[str, Any]:
    """Crude pre-conversion sanity numbers for the source PDF."""
    metrics: Dict[str, Any] = {
        "page_count": 0,
        "char_count": 0,
        "bytes": 0,
    }
    try:
        metrics["bytes"] = os.path.getsize(pdf_path)
    except OSError:
        pass
    try:
        import fitz   # PyMuPDF — already a dep
        with fitz.open(pdf_path) as doc:
            metrics["page_count"] = doc.page_count
            metrics["char_count"] = sum(
                len(p.get_text("text") or "") for p in doc
            )
    except Exception:
        pass
    return metrics


def _measure_converted_docx(docx_path: str) -> Dict[str, Any]:
    """Post-conversion structural numbers used by the score."""
    metrics: Dict[str, Any] = {
        "paragraph_count":      0,
        "non_empty_paragraphs": 0,
        "bullet_paragraphs":    0,
        "section_hits":         0,
        "multi_col_tables":     0,
        "bytes":                0,
        "char_count":           0,
    }
    try:
        metrics["bytes"] = os.path.getsize(docx_path)
    except OSError:
        pass
    try:
        import docx as _docx_lib
        doc = _docx_lib.Document(docx_path)
    except Exception:
        return metrics

    body_chars: list[str] = []
    for p in doc.paragraphs:
        metrics["paragraph_count"] += 1
        text = (p.text or "").strip()
        if text:
            metrics["non_empty_paragraphs"] += 1
            body_chars.append(text)
        try:
            style_name = (p.style.name or "").lower() if p.style else ""
        except Exception:
            style_name = ""
        if style_name.startswith("list") or "bullet" in style_name:
            metrics["bullet_paragraphs"] += 1
            continue
        try:
            if p._p.pPr is not None and p._p.pPr.numPr is not None:
                metrics["bullet_paragraphs"] += 1
        except Exception:
            pass

    for tbl in doc.tables:
        if len(tbl.columns) >= 2:
            metrics["multi_col_tables"] += 1
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = (p.text or "").strip()
                    if text:
                        metrics["non_empty_paragraphs"] += 1
                        body_chars.append(text)

    body_text = "\n".join(body_chars)
    metrics["char_count"] = len(body_text)
    metrics["section_hits"] = len(set(
        m.group(0).lower() for m in _SECTION_RX.finditer(body_text)
    ))
    return metrics


def _compute_convertibility_score(
    pdf_metrics:  Dict[str, Any],
    docx_metrics: Dict[str, Any],
) -> int:
    """
    Combine pre- and post-conversion metrics into a 0-100 score.
    See module docstring for the rubric.
    """
    score = 0
    pdf_chars = pdf_metrics.get("char_count") or 0
    docx_chars = docx_metrics.get("char_count") or 0
    if pdf_chars > 0 and docx_chars >= int(pdf_chars * 0.80):
        # Preserved at least 80% of the original text — strong signal that
        # paragraph extraction worked.
        score += 60
    elif pdf_chars > 0 and docx_chars >= int(pdf_chars * 0.50):
        # Partial preservation — degraded but not catastrophic.
        score += 30

    if docx_metrics.get("section_hits", 0) >= 2:
        score += 20

    if docx_metrics.get("multi_col_tables", 0) == 0:
        score += 10

    pdf_bytes = pdf_metrics.get("bytes") or 0
    docx_bytes = docx_metrics.get("bytes") or 0
    if pdf_bytes > 0 and docx_bytes <= pdf_bytes * 3:
        score += 10

    return max(0, min(100, score))


# ─────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────

def convert_pdf_to_docx(
    pdf_path:    str,
    output_path: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Convert a PDF CV to DOCX using `pdf2docx` and score the conversion.

    Returns:
        {
            "ok":          bool,
            "docx_path":   str,             # only set when ok=True
            "score":       int,             # 0-100
            "threshold":   CONVERTIBILITY_THRESHOLD,
            "acceptable":  bool,            # ok AND score >= threshold
            "pdf_metrics":  {...},          # pre-conversion numbers
            "docx_metrics": {...},          # post-conversion numbers
            "reason":      str,             # only set when ok=False or rejected
        }

    On any failure (missing lib, bad input, conversion crash) returns
    `ok=False` and a descriptive `reason`. The router treats both
    `ok=False` and `acceptable=False` as "fall back to the rebuild path".

    If `output_path` is None, the converted DOCX is written next to the
    input PDF with a `.docx` extension.
    """
    result: Dict[str, Any] = {
        "ok":           False,
        "docx_path":    "",
        "score":        0,
        "threshold":    CONVERTIBILITY_THRESHOLD,
        "acceptable":   False,
        "pdf_metrics":  {},
        "docx_metrics": {},
        "reason":       "",
    }

    if not pdf_path or not os.path.exists(pdf_path):
        result["reason"] = f"PDF not found: {pdf_path!r}"
        return result

    try:
        from pdf2docx import Converter
    except Exception as e:
        result["reason"] = (
            f"pdf2docx unavailable in this build: {type(e).__name__}: {e}"
        )
        return result

    if output_path is None:
        base, _ext = os.path.splitext(pdf_path)
        output_path = base + ".docx"

    # Ensure target dir exists.
    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.isdir(out_dir):
        try:
            os.makedirs(out_dir, exist_ok=True)
        except OSError as e:
            result["reason"] = f"Cannot create output dir: {e}"
            return result

    result["pdf_metrics"] = _measure_pdf(pdf_path)

    # Run the conversion. pdf2docx is chatty by default; we silence it via
    # its own debug flag rather than redirecting stdout (cleaner in tests).
    try:
        cv = Converter(pdf_path)
        # `debug=False` quiets the per-page logs; multi_processing is left
        # at its default (off) — CV PDFs are short and the parallel-page
        # mode adds startup overhead that exceeds the convert time.
        cv.convert(output_path, start=0, end=None)
        cv.close()
    except Exception as e:
        result["reason"] = (
            f"pdf2docx conversion failed: {type(e).__name__}: {e}"
        )
        return result

    # May 2026 fix: deduplicate consecutive paragraphs (pdf2docx bug with
    # table-based layouts that duplicates text 3×).
    _deduplicate_docx_paragraphs(output_path)

    if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
        result["reason"] = "pdf2docx produced no output"
        return result

    result["ok"]           = True
    result["docx_path"]    = output_path
    result["docx_metrics"] = _measure_converted_docx(output_path)
    result["score"]        = _compute_convertibility_score(
        result["pdf_metrics"], result["docx_metrics"],
    )
    result["acceptable"]   = result["score"] >= CONVERTIBILITY_THRESHOLD

    if not result["acceptable"]:
        result["reason"] = (
            f"Convertibility score {result['score']}/100 below threshold "
            f"{CONVERTIBILITY_THRESHOLD} — likely designer/multi-column CV. "
            f"Falling back to rebuild path."
        )

    return result


def _deduplicate_docx_paragraphs(docx_path: str) -> None:
    """
    Remove consecutively duplicated paragraphs (pdf2docx bug).

    May 2026 fix: pdf2docx sometimes duplicates text 3× in table-based PDF
    layouts (e.g., header/contact sections). This deduplication pass clears
    consecutive duplicate paragraphs while preserving unique content.
    """
    try:
        import docx as _docx
    except Exception:
        return  # docx unavailable — skip dedup

    try:
        doc = _docx.Document(docx_path)
    except Exception:
        return  # can't open docx — skip dedup

    seen = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and seen and text == seen[-1]:
            # Clear this duplicate paragraph
            for run in para.runs:
                run.text = ""
        elif text:
            seen.append(text)

    try:
        doc.save(docx_path)
    except Exception:
        pass  # save failed — continue with potentially duplicated docx
