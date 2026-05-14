"""
agents.cv_docx_parser
=====================

DOCX → outline parser (May 2026 / DOCX path).

Produces the same outline shape as `pdf_editor.build_outline` so the
downstream tailor / strategist / reviewer pipeline doesn't know or care
whether the CV came from a PDF or a Word document:

    {
        "summary": "...",
        "roles":   [
            {
                "header":  "Company, Title  Dates",
                "section": "experience" | "projects",
                "bullets": [{"text": "...", "length": int}, ...],
            },
            ...
        ],
        "skills":  ["Python", "React", ...],
    }

Each role and bullet ALSO carries an internal `_anchor` field with the
paragraph index that produced it. `cv_docx_editor.py` uses these anchors
to rewrite text in the original document while preserving all formatting
(font, bold, italic, colour, indent). cv_diff_tailor / tailor_strategist
ignore unknown fields, so the anchors flow through without side-effects.

Why DOCX parsing is structural (not coordinate-based)
------------------------------------------------------
DOCX is a tree of Paragraph → Run elements with explicit style names.
We classify each paragraph using three signals, in order of confidence:

1.  Numbering / list-style → BULLET
2.  Section keyword + heading style or short bold paragraph → SECTION_HEADER
3.  Anything else inside an experience/projects section → ROLE_HEADER

This is far more robust than the PDF path's x-coordinate / font-size
heuristics. The trade-off is that DOCX has no fixed page layout, so we
can't reason about page-breaks — but the tailor doesn't need to, because
text edits in a DOCX never cause overflow in the first place.

Never raises. On any parse failure returns an empty outline; downstream
treats that as "no edits possible" and falls back to the rebuild path.
"""

from __future__ import annotations

import os
import re
from typing import Any, Dict, List, Optional, Tuple


# ─────────────────────────────────────────────────────────────
# Section / heading detection
# ─────────────────────────────────────────────────────────────

# Section header detection. We tolerate four shapes:
#
#   1.  "Summary"                              (whole line is the keyword)
#   2.  "Summary:"                             (keyword followed by colon)
#   3.  "Professional Summary"                 (optional adjective + keyword)
#   4.  "Skills\nProduct: ..."                 (keyword on first line, body
#                                               follows after a line break;
#                                               we check only the first line)
#
# `pdf2docx` output frequently uses shape #3 ("Professional Summary",
# "Professional Experience") and shape #4 (skills inline with categorised
# body). Both are real-world CV writing styles, not pdf2docx artefacts.
_SECTION_RX = re.compile(
    r"^\s*"
    # Optional adjective prefix. "personal" covers "Personal Projects",
    # "Personal Statement". "side" covers "Side Projects".
    r"(?:professional\s+|technical\s+|work\s+|core\s+|featured\s+|"
    r"academic\s+|personal\s+|side\s+|relevant\s+|key\s+|selected\s+)?"
    r"(summary|profile|objective|about(?:\s+me)?|statement|"
    r"experience|employment(?:\s*history)?|history|"
    r"education|achievements?|background|"
    r"skills?|competenc(?:y|ies)|expertise|"
    r"projects?|portfolio|"
    r"certifications?|awards?|publications?|languages?)"
    r"\s*:?\s*$",
    re.IGNORECASE,
)

# Max length for a paragraph to even be considered as a section header.
# A real header is short ("Skills", "Professional Experience"); a long
# paragraph that happens to mention "experience" is prose, not a section
# break.
_MAX_SECTION_HEADER_CHARS = 60

# Map matched section keyword (group 1 of `_SECTION_RX`) → canonical
# section type used by the downstream pipeline. `pdf_editor.build_outline`
# emits these same labels.
#
# Note: the regex's optional prefix ("professional ", "technical ", etc.)
# is stripped from group 1, so we map only the bare keyword.
_SECTION_CANONICAL: Dict[str, str] = {
    "summary":         "summary",
    "profile":         "summary",
    "objective":       "summary",
    "about":           "summary",
    "about me":        "summary",
    "statement":       "summary",
    "experience":      "experience",
    "employment":      "experience",
    "employment history": "experience",
    "history":         "experience",
    "projects":        "projects",
    "project":         "projects",
    "portfolio":       "projects",
    "skills":          "skills",
    "skill":           "skills",
    "competency":      "skills",
    "competencies":    "skills",
    "expertise":       "skills",
    # "achievements" and "background" only count as a section when the
    # paragraph also had an "academic " prefix (regex catches both with
    # the same group). Map both to education.
    "education":       "education",
    "achievements":    "education",
    "background":      "education",
    "certifications":  "certifications",
    "certification":   "certifications",
    "awards":          "awards",
    "award":           "awards",
    "publications":    "publications",
    "publication":     "publications",
    "languages":       "languages",
    "language":        "languages",
}

# A paragraph is a "list / bullet" paragraph when any of these hold.
# Conservative — false negatives only mean we treat the bullet as a role
# header (the editor will then refuse to rewrite it, which is safe).
_BULLET_STYLE_HINTS = ("list bullet", "list paragraph", "bullet", "list number")

# Date pattern in role headers ("Jan 2024 – Present", "2021-2023").
# Used to distinguish role headers from prose paragraphs when the
# CV doesn't use heading styles.
_DATE_HINT_RX = re.compile(
    r"(?:\d{4}|"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*"
    r"\s*\.?\s*\d{2,4})",
    re.IGNORECASE,
)

# Bullet glyphs that may appear inside paragraph text (some templates
# embed the bullet character rather than using list numbering).
# Trailing whitespace is OPTIONAL: pdf2docx output produces `•Started…`
# with no space; native Word output usually has `• Started…`. Both must
# match. We require either trailing whitespace OR a non-alphanumeric
# follower (so we don't false-match a hyphen in "data-driven").
_BULLET_GLYPH_RX = re.compile(
    r"^\s*[\u2022\u00b7\u25aa\u25cb\u25a0\u2043\u2219\u25b8\u25b6]\s*"
    r"|"
    r"^\s*[\-\*]\s+",
)


# ─────────────────────────────────────────────────────────────
# Paragraph classification helpers
# ─────────────────────────────────────────────────────────────

def _is_bullet_paragraph(paragraph: Any) -> bool:
    """True if this paragraph is rendered as a bulleted/numbered list item."""
    # 1. Style name hints (most reliable for Word-generated DOCX).
    try:
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    except Exception:
        style_name = ""
    if any(hint in style_name for hint in _BULLET_STYLE_HINTS):
        return True

    # 2. Explicit numbering definition in the paragraph properties.
    try:
        ppr = paragraph._p.pPr
        if ppr is not None and ppr.numPr is not None:
            return True
    except Exception:
        pass

    # 3. Leading bullet glyph in the text itself (templates that "fake" bullets).
    text = (paragraph.text or "").strip()
    if text and _BULLET_GLYPH_RX.match(text):
        return True

    return False


def _strip_leading_bullet_glyph(text: str) -> str:
    """If text starts with a bullet glyph, drop it (and the following space)."""
    return _BULLET_GLYPH_RX.sub("", text, count=1).strip()


# Internal bullet split pattern. Detects when pdf2docx has collapsed multiple
# bullets into a single paragraph by looking for newline + bullet-glyph
# sequences inside the paragraph text. Run 19 evidence: Rishav's PDF gave
# pdf2docx output where 4-6 separate bullets ended up concatenated as one
# 500-700 char "bullet" with internal `\n• ...\n• ...` structure.
_INTERNAL_BULLET_SPLIT_RX = re.compile(
    r"\n\s*[•·▪○■⁃∙▸▶\-\*]\s+"
)


def _split_mega_bullet(text: str) -> List[str]:
    """
    Split a paragraph that pdf2docx concatenated from multiple bullets back
    into atomic bullets.

    Returns a list of bullet texts (with leading glyphs stripped). When the
    paragraph contains no internal bullet boundaries the returned list is
    `[text]` (single element).

    Splits ONLY when:
    - The paragraph contains at least one internal `\\n + bullet-glyph`
      pattern (very strong signal pdf2docx fused bullets)
    OR
    - The paragraph is >250 chars AND contains 2+ internal newlines that
      look like wrap boundaries (fallback heuristic for when pdf2docx
      stripped glyphs but kept paragraph breaks)
    """
    if not text:
        return []
    stripped = text.strip()
    if not stripped:
        return []

    # Primary split: internal bullet glyphs after newline.
    parts = _INTERNAL_BULLET_SPLIT_RX.split(stripped)
    if len(parts) > 1:
        return [
            _strip_leading_bullet_glyph(p).strip()
            for p in parts
            if p.strip()
        ]

    # Fallback split A: long paragraph with internal newlines where each
    # line starts like a bullet (Capitalised verb + space + content). This
    # catches pdf2docx outputs where bullet glyphs were stripped but the
    # original line breaks were preserved.
    if len(stripped) > 250 and "\n" in stripped:
        # Common CV bullet-starter verbs (past tense). Casing matters.
        _BULLET_VERB_RX = re.compile(
            r"^(?:Built|Led|Drove|Delivered|Managed|Owned|Designed|"
            r"Developed|Implemented|Created|Authored|Defined|Established|"
            r"Architected|Engineered|Shipped|Launched|Coordinated|"
            r"Orchestrated|Spearheaded|Partnered|Collaborated|Improved|"
            r"Reduced|Increased|Achieved|Generated|Identified|Resolved|"
            r"Negotiated|Facilitated|Conducted|Analysed|Analyzed|"
            r"Streamlined|Translated|Drafted|Wrote|Researched|Initiated)\b"
        )
        line_split = re.split(r"\s*\n\s*", stripped)
        if len(line_split) >= 2:
            verb_starters = sum(
                1 for ln in line_split if _BULLET_VERB_RX.match(ln.strip())
            )
            # If ≥50% of the lines start with a bullet-verb, treat as a
            # merged-bullet paragraph and split.
            if verb_starters >= max(2, len(line_split) // 2):
                out: List[str] = []
                for chunk in line_split:
                    chunk = chunk.strip()
                    if not chunk:
                        continue
                    out.append(_strip_leading_bullet_glyph(chunk))
                if len(out) > 1:
                    return out

    # Fallback split B: long paragraph with sentence-terminated internal
    # newlines. Catches paragraphs where each bullet was a single sentence
    # ending with a period and pdf2docx kept them on separate lines.
    if len(stripped) > 250 and stripped.count("\n") >= 2:
        sentence_split = re.split(r"\.\s*\n\s*", stripped)
        if len(sentence_split) >= 2:
            out: List[str] = []
            for i, chunk in enumerate(sentence_split):
                chunk = chunk.strip()
                if not chunk:
                    continue
                if i < len(sentence_split) - 1 and not chunk.endswith("."):
                    chunk = chunk + "."
                out.append(_strip_leading_bullet_glyph(chunk))
            if len(out) > 1:
                return out

    return [_strip_leading_bullet_glyph(stripped)]


def _is_section_header(text: str) -> Optional[str]:
    """
    Return the canonical section type ("summary"/"experience"/...) if `text`
    looks like a section header line. Otherwise None.

    Match rules:
      - Only consider paragraphs whose first non-empty line is short
        (≤_MAX_SECTION_HEADER_CHARS) — otherwise it's prose that happens
        to contain a section keyword.
      - The first line must MATCH _SECTION_RX (whole line is the keyword,
        optionally with a "professional " / "technical " / etc. prefix
        and a trailing colon).
      - Multi-line paragraphs are checked on the first line only. This
        handles pdf2docx's "Skills\nProduct: ..." pattern where the
        section keyword sits on its own line.
    """
    if not text:
        return None
    # First non-empty line only.
    first_line = ""
    for line in text.splitlines():
        s = line.strip()
        if s:
            first_line = s
            break
    if not first_line or len(first_line) > _MAX_SECTION_HEADER_CHARS:
        return None
    m = _SECTION_RX.match(first_line)
    if not m:
        return None
    keyword = m.group(1).lower().strip()
    return _SECTION_CANONICAL.get(keyword)


def _paragraph_is_bold_heading(paragraph: Any) -> bool:
    """
    Heuristic: paragraph "looks like" a role header even without using a
    Heading style. True when every non-empty run is bold AND the text is
    relatively short (< 200 chars) AND no leading bullet glyph.

    Some CVs (especially conversions from PDFs) emit role headers as
    bold paragraphs styled "Normal". This catches them.
    """
    text = (paragraph.text or "").strip()
    if not text or len(text) > 200:
        return False
    runs = list(paragraph.runs) if hasattr(paragraph, "runs") else []
    if not runs:
        return False
    bold_chars = 0
    total_chars = 0
    for run in runs:
        rt = (run.text or "")
        if not rt.strip():
            continue
        total_chars += len(rt.strip())
        if run.bold:
            bold_chars += len(rt.strip())
    if total_chars == 0:
        return False
    return (bold_chars / total_chars) >= 0.8


# ─────────────────────────────────────────────────────────────
# Paragraph walk — flatten body + tables in document order
# ─────────────────────────────────────────────────────────────

def _iter_body_paragraphs(doc: Any) -> List[Tuple[int, Any]]:
    """
    Return `[(global_index, paragraph), ...]` in TRUE document order
    (body paragraphs and table-cell paragraphs interleaved as they
    appear in the document tree).

    Why this matters: `pdf2docx` reconstructs CV layout using tables for
    multi-column blocks (a one-row-by-two-column table is how it models
    a "role header LEFT | dates RIGHT" line). If we walked `doc.paragraphs`
    first and `doc.tables` second (the obvious python-docx approach), the
    table content would clump at the end of the parse out of order with
    the prose around it — every role header inside a table would land
    AFTER every body paragraph, breaking the section state machine.

    The fix is to walk `doc.element.body` in raw element order and yield
    paragraphs / table-cell-paragraphs as they appear. The global_index
    is a stable monotonic id the editor uses to relocate the same
    paragraph on re-open.

    NOTE: the editor's `_index_paragraphs` MUST mirror this walk shape
    exactly. Both functions live in the same module pair on purpose.
    """
    from docx.oxml.ns import qn   # late import — avoids hard dep at module load

    out: List[Tuple[int, Any]] = []
    idx = 0

    # Late-import lookups so `cv_docx_parser` stays importable even when
    # docx is missing (the build_outline call already guards that).
    try:
        from docx.text.paragraph import Paragraph as _Paragraph
        from docx.table import Table as _Table, _Cell
    except Exception:
        # Fallback: pre-1.0 python-docx layout. The walk below still works
        # via element tags, just without typed wrappers.
        _Paragraph = _Table = _Cell = None   # type: ignore

    P_TAG  = qn("w:p")
    TBL_TAG = qn("w:tbl")

    def _walk_block_container(parent_element: Any, container_obj: Any) -> None:
        nonlocal idx
        for child in parent_element.iterchildren():
            tag = child.tag
            if tag == P_TAG:
                if _Paragraph is not None:
                    p = _Paragraph(child, container_obj)
                else:
                    # Last-resort: use python-docx's known paragraph proxy.
                    p = doc.paragraphs[0].__class__(child, container_obj)
                out.append((idx, p))
                idx += 1
            elif tag == TBL_TAG:
                if _Table is not None:
                    tbl = _Table(child, container_obj)
                    for row in tbl.rows:
                        for cell in row.cells:
                            # Walk the cell's element to recover paragraph
                            # AND nested table order. CVs rarely nest more
                            # than one level deep; this handles two.
                            _walk_block_container(cell._tc, cell)

    _walk_block_container(doc.element.body, doc)
    return out


# ─────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────

def build_outline_from_docx(docx_path: str) -> Dict[str, Any]:
    """
    Parse a DOCX file and return an outline matching the shape of
    `agents.pdf_editor.build_outline`. Never raises — returns an empty
    outline on any failure so the caller can fall back to the rebuild path.

    `roles[i]["_anchor"]` and `roles[i]["bullets"][j]["_anchor"]` contain
    the paragraph indices from `_iter_body_paragraphs`. `cv_docx_editor`
    uses these to relocate paragraphs in the on-disk DOCX without
    re-parsing the outline.
    """
    # `_summary_anchors` and `_skills_anchors` are internal fields the
    # editor uses to relocate the summary / skills paragraphs without
    # re-parsing. cv_diff_tailor and tailor_strategist only read
    # `summary` / `roles` / `skills`, so unknown fields pass through harmlessly.
    out: Dict[str, Any] = {
        "summary": "", "roles": [], "skills": [],
        "_summary_anchors": [], "_skills_anchors": [],
    }

    if not docx_path or not os.path.exists(docx_path):
        return out

    try:
        import docx as _docx_lib
        doc = _docx_lib.Document(docx_path)
    except Exception:
        return out

    paragraphs = _iter_body_paragraphs(doc)
    if not paragraphs:
        return out

    # Phase 1: classify every paragraph.
    #
    # Each entry: {
    #   "idx": int (anchor),
    #   "text": str,
    #   "kind": "section" | "bullet" | "header_like" | "prose",
    #   "section": Optional[str]  (only set for kind="section"),
    # }
    classified: List[Dict[str, Any]] = []
    for idx, p in paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        section = _is_section_header(text)
        if section:
            classified.append({
                "idx": idx, "text": text,
                "kind": "section", "section": section,
            })
            continue
        if _is_bullet_paragraph(p):
            # Run 19 audit fix: detect and split pdf2docx-merged megabullets.
            # When pdf2docx concatenates 4-6 separate bullets into a single
            # paragraph (typical for PDF→DOCX round-trips), we need to split
            # them back into atomic bullets so the tailor can rewrite each one
            # without the length guard rejecting (rewrite would be way shorter
            # than the merged original).
            sub_bullets = _split_mega_bullet(text)
            if len(sub_bullets) > 1:
                print(
                    f"   ✂️  cv_docx_parser: split megabullet at anchor={idx} "
                    f"into {len(sub_bullets)} atomic bullets "
                    f"(orig_len={len(text)})"
                )
                for sub_idx, sub_text in enumerate(sub_bullets):
                    if not sub_text:
                        continue
                    classified.append({
                        "idx":     idx,
                        "sub_idx": sub_idx,
                        "sub_count": len(sub_bullets),
                        "text":    sub_text,
                        "kind":    "bullet",
                        "section": None,
                    })
                continue
            classified.append({
                "idx": idx,
                "text": _strip_leading_bullet_glyph(text),
                "kind": "bullet", "section": None,
            })
            continue
        if _paragraph_is_bold_heading(p):
            classified.append({
                "idx": idx, "text": text,
                "kind": "header_like", "section": None,
            })
            continue
        classified.append({
            "idx": idx, "text": text,
            "kind": "prose", "section": None,
        })

    # Phase 2: walk classified stream, grouping by section.
    #
    # State machine:
    #   - On "section": flush any in-progress role; switch current section.
    #   - In "summary"/"profile" section: accumulate prose lines into summary.
    #   - In "experience"/"projects" section: header_like / prose-with-date
    #     starts a new role; subsequent bullets attach to that role.
    #   - In "skills" section: accumulate prose into a skills text blob,
    #     then split on common delimiters.
    current_section: str = "preamble"   # before any section header
    current_role: Optional[Dict[str, Any]] = None
    summary_chunks: List[str] = []
    skills_chunks: List[str] = []

    def _flush_role() -> None:
        nonlocal current_role
        if current_role is not None and current_role.get("bullets"):
            out["roles"].append(current_role)
        current_role = None

    # Pre-section fallback: the first prose paragraph(s) before any explicit
    # section header are often a Summary block (templates that skip the
    # "Summary" label). Capture them into a pre-section buffer; if the CV
    # never has an explicit Summary header, we'll use them.
    pre_section_prose: List[Tuple[int, str]] = []   # (anchor, text)

    # Refactored state machine (May 2026 audit follow-up). pdf2docx output
    # produces three artefacts the naive walk above doesn't handle:
    #
    #   (a) ROLE FRAGMENTS — a single role header gets split across multiple
    #       paragraphs (company name on one line, dates on the next, client
    #       on the third). Treat consecutive header_like / short-prose
    #       paragraphs as fragments of ONE header until a bullet shows up.
    #
    #   (b) BULLET CONTINUATIONS — when a bullet's text wraps, pdf2docx
    #       emits the continuation as a plain (unstyled) prose paragraph.
    #       Append it to the most recent bullet rather than treating it
    #       as a new header.
    #
    #   (c) EMPTY BULLETS — pdf2docx sometimes emits a bullet-styled
    #       paragraph with no text (the cell that holds just "•"). Skip.
    in_bullet_streak = False   # True after we've seen a bullet in this role
    for entry in classified:
        kind = entry["kind"]
        text = entry["text"]
        idx  = entry["idx"]

        if kind == "section":
            _flush_role()
            current_section = entry["section"] or "other"
            in_bullet_streak = False
            continue

        if current_section == "preamble":
            # Skip name/email-style top-of-CV header lines (short, often
            # contain "@" or phone numbers). Save longer prose as a
            # possible summary candidate when the CV omits a header.
            if len(text) > 40 and "@" not in text and not re.search(r"\+?\d[\d\s().-]{6,}", text):
                pre_section_prose.append((idx, text))
            continue

        if current_section == "summary":
            if kind in ("prose", "header_like", "bullet"):
                summary_chunks.append(text)
                out["_summary_anchors"].append(idx)
            continue

        if current_section in ("experience", "projects"):
            if kind == "bullet":
                # Skip empty bullets — pdf2docx artifacts from blank
                # table cells used as layout separators.
                if not text:
                    continue
                if current_role is None:
                    # Bullet without a header — synthesise a placeholder.
                    current_role = {
                        "header":  "(role)",
                        "section": current_section,
                        "bullets": [],
                        "_anchor": idx,
                    }
                bullet_entry = {
                    "text":          text,
                    "length":        len(text),
                    "_anchor":       idx,
                    # `_continuation_anchors` lists the paragraph indices
                    # whose text was merged into this bullet by the
                    # wrap-line continuation logic below.
                    "_continuation_anchors": [],
                }
                # Run 19 audit fix: carry megabullet sub-index forward so
                # the editor can group multi-bullet rewrites back into a
                # single paragraph at apply time.
                if "sub_idx" in entry:
                    bullet_entry["_megabullet_subidx"] = entry["sub_idx"]
                    bullet_entry["_megabullet_count"] = entry.get("sub_count", 1)
                current_role["bullets"].append(bullet_entry)
                in_bullet_streak = True
                continue

            # Non-bullet content (header_like / prose). Interpretation
            # depends on whether we're mid-bullet-streak.

            if not in_bullet_streak:
                # Pre-bullets phase: this is part of the role header (or
                # the role hasn't started yet). Merge into current role's
                # header — these are role-header FRAGMENTS, not bullets.
                if current_role is None:
                    current_role = {
                        "header":  text,
                        "section": current_section,
                        "bullets": [],
                        "_anchor": idx,
                    }
                else:
                    current_role["header"] = (
                        current_role["header"].rstrip() + " " + text
                    ).strip()
                continue

            # Mid-bullet-streak. Conservative rule: only flush + start a
            # new role when the paragraph is STRONGLY signalled as a new
            # role header. Signals are bold-heavy formatting OR a date
            # hint (e.g. "Aug 2023 – Jan 2024"). Plain prose paragraphs
            # are bullet-text continuations from pdf2docx text wrap, not
            # new headers. Anything else gets appended to the most recent
            # bullet's text so we don't lose CV content.
            is_strong_new_role = (
                kind == "header_like"
                or _DATE_HINT_RX.search(text) is not None
            )
            if is_strong_new_role:
                _flush_role()
                current_role = {
                    "header":  text,
                    "section": current_section,
                    "bullets": [],
                    "_anchor": idx,
                }
                in_bullet_streak = False
            else:
                if current_role and current_role["bullets"]:
                    last = current_role["bullets"][-1]
                    last["text"]   = (last["text"].rstrip() + " " + text).strip()
                    last["length"] = len(last["text"])
                    last["_continuation_anchors"].append(idx)
            continue

        if current_section == "skills":
            if kind in ("prose", "bullet", "header_like"):
                skills_chunks.append(text)
                out["_skills_anchors"].append(idx)
            continue

        # Other sections (education / certifications / etc.) are recorded
        # but not surfaced — the tailor doesn't edit them in v1.

    _flush_role()

    # Summary fallback: if we never hit an explicit Summary section but
    # we did capture some pre-section prose, use that.
    if not summary_chunks and pre_section_prose:
        summary_chunks = [text for _idx, text in pre_section_prose]
        out["_summary_anchors"] = [idx for idx, _text in pre_section_prose]

    out["summary"] = " ".join(s.strip() for s in summary_chunks).strip()

    # Skills: flatten and split on common delimiters. If the skills text
    # looks "categorised" (contains "Foo: bar, baz" pattern), keep it as
    # a single string so the tailor knows not to reorder — same policy as
    # pdf_editor.build_outline.
    skills_text = " ".join(s.strip() for s in skills_chunks).strip()
    if skills_text:
        if re.search(r"\b[A-Z][A-Za-z &/]{2,20}:\s", skills_text):
            out["skills"] = []   # categorised — reorder disabled
        else:
            items = [
                s.strip() for s in re.split(r"[,;|\u2022\u00b7]", skills_text)
                if s.strip()
            ]
            out["skills"] = items

    return out


def outline_anchors(outline: Dict[str, Any]) -> List[Tuple[str, int]]:
    """
    Convenience helper: flatten an outline into a list of
    `(kind, paragraph_index)` pairs. Used by `cv_docx_editor` to validate
    that the document hasn't shifted between parse and edit.

    Kinds: "role_header", "bullet". Summary anchors are intentionally
    excluded because the summary is treated as a contiguous block.
    """
    pairs: List[Tuple[str, int]] = []
    for role in outline.get("roles", []):
        anc = role.get("_anchor")
        if isinstance(anc, int):
            pairs.append(("role_header", anc))
        for b in role.get("bullets", []):
            banc = b.get("_anchor") if isinstance(b, dict) else None
            if isinstance(banc, int):
                pairs.append(("bullet", banc))
    return pairs


def _outline_quality_ok(outline: Dict[str, Any]) -> bool:
    """
    Decide whether the parsed outline is usable by the DOCX editor.

    Rewritten May 2026 (run 17 follow-up): the previous check rejected
    outlines whose role headers were long, newline-heavy, or "garbled".
    That caught true pdf2docx corruption (3x duplicated text) but ALSO
    caught legitimate table-based CVs where the right column contains
    descriptive "Project: ..." labels that pdf2docx faithfully preserves
    as long paragraphs. Those outlines are functionally fine — the
    bullets are correctly anchored to their paragraphs, and the editor
    rewrites bullets, not headers. Rejecting them sent the run to the
    PyMuPDF in-place fallback which shrinks fonts on every rewrite.

    The new check is purely functional: an outline is usable if at
    least one role has at least one bullet that the editor can rewrite.
    Header cosmetics (length, newlines, "Project: ..." prefixes) are
    irrelevant because the header is just a lookup key — the actual
    edit lands on the bullet's paragraph anchor.

    Returns False only when the outline is structurally unusable:
      - No roles at all (pdf2docx produced an empty document or the
        parser couldn't identify any experience/projects section)
      - Every role has zero anchored bullets (no editable content)
    """
    roles = outline.get("roles") or []
    if not roles:
        # Run 18 audit: log explicit reason so we can diagnose DOCX
        # rejections in production instead of guessing.
        print(
            "   🔍 _outline_quality_ok: FAIL — outline has 0 roles "
            "(pdf2docx likely produced unstructured text; parser couldn't "
            "identify any experience/projects section)"
        )
        return False

    total_anchored_bullets = 0
    roles_with_zero_bullets = 0
    for role in roles:
        role_bullets = role.get("bullets") or []
        if not role_bullets:
            roles_with_zero_bullets += 1
            continue
        for b in role_bullets:
            if isinstance(b, dict) and isinstance(b.get("_anchor"), int):
                total_anchored_bullets += 1
            elif isinstance(b, dict) and b.get("text"):
                total_anchored_bullets += 1
    if total_anchored_bullets == 0:
        print(
            f"   🔍 _outline_quality_ok: FAIL — {len(roles)} role(s) found "
            f"but 0 anchored bullets ({roles_with_zero_bullets} role(s) with "
            f"empty bullet lists). pdf2docx likely failed to preserve bullet "
            f"glyphs or list-paragraph styling. Sample role headers: "
            f"{[r.get('header','')[:60] for r in roles[:3]]!r}"
        )
        return False

    print(
        f"   ✅ _outline_quality_ok: PASS — {len(roles)} role(s), "
        f"{total_anchored_bullets} anchored bullet(s)"
    )
    return True
